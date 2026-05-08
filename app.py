import streamlit as st
import cv2
import numpy as np
import pytesseract
from pytesseract import Output
from docx import Document
import io
from PIL import Image

# ── Page config ──────────────────────────────────────────────
st.set_page_config(
    page_title="Image to Word Converter",
    page_icon="🖼️",
    layout="centered"
)

st.title("🖼️ Image to Word Converter")
st.write("Upload an image containing text and download it as an editable Word document.")

# ── Helper functions (inline so no local imports needed) ─────

def preprocess_image(pil_image):
    img = np.array(pil_image.convert("RGB"))
    img_bgr = cv2.cvtColor(img, cv2.COLOR_RGB2BGR)
    gray = cv2.cvtColor(img_bgr, cv2.COLOR_BGR2GRAY)
    thresh = cv2.threshold(gray, 150, 255, cv2.THRESH_BINARY)[1]
    return thresh

def extract_text(image):
    data = pytesseract.image_to_data(image, output_type=Output.DICT)
    return data

def detect_formatting(data):
    formatted_words = []
    heights = [h for h in data['height'] if h > 0]
    if not heights:
        return formatted_words
    avg_height = sum(heights) / len(heights)
    for i, text in enumerate(data['text']):
        if text.strip() == "":
            continue
        bold = data['height'][i] > avg_height * 1.15
        formatted_words.append({
            'text': text + ' ',
            'bold': bold,
            'italic': False
        })
    return formatted_words

def generate_doc(words):
    doc = Document()
    para = doc.add_paragraph()
    for word in words:
        run = para.add_run(word['text'])
        run.bold = word['bold']
        run.italic = word['italic']
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

# ── UI ───────────────────────────────────────────────────────

uploaded_file = st.file_uploader("Choose an image", type=["png", "jpg", "jpeg"])

if uploaded_file:
    pil_image = Image.open(uploaded_file)
    st.image(pil_image, caption="Uploaded Image", use_container_width=True)

    with st.spinner("Extracting text and generating Word document..."):
        processed = preprocess_image(pil_image)
        text_data = extract_text(processed)
        words = detect_formatting(text_data)
        doc_buffer = generate_doc(words)

    if words:
        st.success(f"✅ Done! Extracted {len(words)} words.")
        st.download_button(
            label="📥 Download Word Document",
            data=doc_buffer,
            file_name="output.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        with st.expander("Preview extracted text"):
            preview = " ".join(w['text'] for w in words)
            st.write(preview)
    else:
        st.warning("No text could be extracted from this image. Try a clearer image.")
