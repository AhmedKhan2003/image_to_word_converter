from docx import Document


def generate_doc(words, output_path="output.docx"):
    doc = Document()
    para = doc.add_paragraph()

    for word in words:
        run = para.add_run(word['text'])
        run.bold = word['bold']
        run.italic = word['italic']

    doc.save(output_path)
